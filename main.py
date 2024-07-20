from docx import Document
import os
import re
from InvalidFileTypeError import InvalidFileTypeError

ORIGINAL_DOCX_FILES_DIRECTORY = 'original_docx_files'
UPDATED_DOCX_FILES_DIRECTORY = 'updated_docx_files'


def parse_file_to_process():
    file_name = input("File name: ")

    current_path = os.path.dirname(os.path.realpath(__file__))
    directory_path = os.path.join(current_path, ORIGINAL_DOCX_FILES_DIRECTORY)
    file_path = os.path.join(directory_path, file_name)

    if not os.path.isfile(file_path):
        raise FileNotFoundError(f'docx file does not exist in the directory {directory_path}')

    _, file_extension = os.path.splitext(file_path)

    if file_extension.lower() != '.docx':
        raise InvalidFileTypeError(file_extension, '.docx')

    return file_name, file_path


def parse_updated_file():
    file_name = input("File name: ")

    current_path = os.path.dirname(os.path.realpath(__file__))
    directory_path = os.path.join(current_path, UPDATED_DOCX_FILES_DIRECTORY)
    file_path = os.path.join(directory_path, file_name)

    _, file_extension = os.path.splitext(file_path)

    if file_extension.lower() != '.docx':
        raise InvalidFileTypeError(file_extension, '.docx')

    return file_path


def validate_directories():
    current_path = os.path.dirname(os.path.realpath(__file__))
    original_docx_files_path = os.path.join(current_path, 'original_docx_files')
    updated_docx_files_path = os.path.join(current_path, 'updated_docx_files')

    if not os.path.isdir(original_docx_files_path):
        raise NotADirectoryError(f"original_docx_files directory is not found within {current_path}")

    if not os.path.isdir(updated_docx_files_path):
        raise NotADirectoryError(f"updated_docx_files directory is not found within {current_path}")


def extract_keywords(document):
    keywords = set()
    pattern = re.compile(r'\[(.*?)]')

    for para in document.paragraphs:
        keywords.update(pattern.findall(para.text.replace('\n', ' ').strip()))

    return keywords


def parse_replacements(keyword):
    print(f'What is the replacement for keyword: {keyword}?')
    return input('Enter the replacement: ')


def get_user_replacements(keywords):
    replacements = {}
    for keyword in keywords:
        replacement = parse_replacements(keyword)
        replacements[keyword] = replacement
    return replacements


def replace_keywords(text, replacements):
    pattern = re.compile(r'\[(.*?)]')

    def replacer(match):
        keyword = match.group(1)
        return replacements.get(keyword, match.group(0))

    return pattern.sub(replacer, text)


def save_replaced_text_with_formatting(document, file_path, replacements):
    for para in document.paragraphs:
        for run in para.runs:
            run.text = replace_keywords(run.text, replacements)

    try:
        document.save(file_path)
    except PermissionError as e:
        print(f'Permission Error{e}\nProbably left word document open')


def main():
    validate_directories()
    print("Please input the name of the file you'd like to process")
    file_name, file_path = parse_file_to_process()
    document = Document(file_path)
    keywords = extract_keywords(document)
    replacements = get_user_replacements(keywords)
    print("Please input the name of the new file")
    new_file_path = parse_updated_file()
    save_replaced_text_with_formatting(document, new_file_path, replacements)


if __name__ == "__main__":
    main()
